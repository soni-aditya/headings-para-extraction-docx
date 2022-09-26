# python-docx package
from docx import Document

def readtxt(filename):
    doc = Document(filename)
    ext_para = []
    ext_head = []
    other = []
    for para in doc.paragraphs:
        if "Normal" in para.style.name:
            ext_para.append(para.text)
        elif "Heading" in para.style.name:
            ext_head.append(para.text)
        else:
            other.append(para.text)
    return ext_para, ext_head, other

# doc = Document('sample2.docx')
# print(len(doc.paragraphs))
#
# print(doc.paragraphs.text)


ext_para, ext_head, other = readtxt('sample2.docx')
print("HEADINGS ----------------")
print(ext_head)
print("PARAS ----------------------------")
print(ext_para)

